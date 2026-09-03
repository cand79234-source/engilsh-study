"""新功能测试：导入弹窗上传 docx 文件 → 提取回填 → 解析导入 → 学习页可见。

原来这里硬编码了一份临时上传的 docx 绝对路径，那个文件早已不存在，
测试一跑就 FileNotFoundError。现在改为运行时自动生成一份等价 fixture
（同样是「第2周｜工作与日常｜120词」+ 6 个 Day 分组 + 110 个词），
测试自包含，不依赖任何外部文件。
"""
import os, sys
os.environ["EOS_DB"] = "/tmp/eos_file_upload_test.db"
if os.path.exists("/tmp/eos_file_upload_test.db"):
    os.remove("/tmp/eos_file_upload_test.db")
sys.path.insert(0, "backend")

from playwright.sync_api import sync_playwright

BASE = "http://127.0.0.1:8000"
DOCX = "/tmp/eos_import_fixture.docx"

# Day 名称固定 6 天，词数合计 110（Day1~Day5 各 18，Day6 20）
_DAY_NAMES = ["第一天上班", "同事与会议", "项目与进度", "邮件与汇报",
              "客户与谈判", "周末与生活"]
_WORDS = """
colleague manager schedule deadline meeting agenda report project budget client
contract invoice proposal feedback summary progress deliver assign approve review
revise submit confirm postpone reschedule attend prepare present discuss negotiate
clarify update follow remind arrange handle solve improve reduce increase achieve
target result issue risk plan strategy priority resource workload overtime
promotion salary bonus benefit training skill experience position department
branch headquarters supplier partner competitor market customer service quality
standard process system platform tool document template version backup access
permission password account profile setting notice policy rule guideline
procedure record archive folder attachment reply forward inbox outbox draft
signature calendar reminder invitation venue lunch coffee weekend holiday
""".split()


def _ensure_fixture():
    """生成一份与用户真实 docx 结构等价的导入文件。"""
    if os.path.exists(DOCX):
        os.remove(DOCX)
    from docx import Document
    doc = Document()
    doc.add_paragraph("第2周｜工作与日常｜120词")
    idx = 0
    for d, name in enumerate(_DAY_NAMES, start=1):
        doc.add_paragraph(f"Day {d}｜{name}")
        n = 20 if d == 6 else 18
        for _ in range(n):
            w = _WORDS[idx % len(_WORDS)]
            idx += 1
            doc.add_paragraph(f"{w} — 释义{idx}")
            doc.add_paragraph(f"- This is a sentence about {w}.")
            doc.add_paragraph(f"这是关于{w}的句子。")
    doc.save(DOCX)


_ensure_fixture()

passed = failed = 0
def check(name, cond, extra=""):
    global passed, failed
    if cond:
        passed += 1
        print(f"  ✅ {name}")
    else:
        failed += 1
        print(f"  ❌ {name} {extra}")

with sync_playwright() as pw:
    browser = pw.chromium.launch()
    page = browser.new_page(viewport={"width": 414, "height": 896})  # 手机视口
    errors = []
    page.on("pageerror", lambda e: errors.append(str(e)))

    print("=== 1. 打开学习页，进入导入弹窗 ===")
    page.goto(BASE, wait_until="networkidle")
    # 找到导入入口（顶部工具区「导入」按钮）
    page.click("text=导入")
    page.wait_for_selector("#imp_rich", timeout=5000)
    check("导入弹窗打开", page.is_visible("#imp_rich"))
    check("文件上传按钮存在", page.is_visible("#imp_file_btn"))
    check("隐藏文件选择框存在", page.locator("#imp_file").count() == 1)

    print("=== 2. 上传用户真实 docx（写作.docx） ===")
    page.set_input_files("#imp_file", DOCX)
    page.wait_for_selector("#imp_file_result .res", timeout=15000)
    res_html = page.inner_html("#imp_file_result")
    check("显示提取成功提示", "已读取" in res_html, res_html[:100])
    check("提示格式与行数", "docx" in res_html and "行" in res_html)
    ta = page.input_value("#imp_rich")
    check("文本已回填 textarea", "第2周｜工作与日常｜120词" in ta)
    check("回填内容含 Day 1 分组", "Day 1｜第一天上班" in ta)
    check("零宽字符已清洗", "\ufeff" not in ta)

    print("=== 3. 点「解析并导入」 ===")
    page.click("button:has-text('解析并导入')")
    page.wait_for_selector("#imp_rich_result .res", timeout=20000)
    r = page.inner_html("#imp_rich_result")
    check("导入成功提示", "已合并到 第2周" in r, r[:150])
    check("词数统计正确", "110 词" in r, r[:200])
    check("天数统计正确", "6 天" in r)

    print("=== 4. 无 JS 错误 ===")
    check("无JS错误", not errors, str(errors[:3]))

    browser.close()

print(f"\n========== 结果: {passed} 通过, {failed} 失败 ==========")
sys.exit(1 if failed else 0)
